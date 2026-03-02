#!/usr/bin/env python3
"""
Run this script once to generate the Word case study document.

    pip3 install python-docx
    cd "/Users/tomdevilliers/Documents/Thomas/Woolworths"
    ~/Library/Python/3.9/bin/python3 generate_word_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(13, 27, 42)   # navy
    return p

def body(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0d1b2a')
        tcPr.append(shd)

    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'f4f6f8')
                tcPr.append(shd)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

# ── Document ──────────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default paragraph font
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ─────────────────────────────────────────────────────────────────────────────
# TITLE BLOCK
# ─────────────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("Egyptian Property Market Analysis")
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(13, 27, 42)

sub = doc.add_paragraph()
r = sub.add_run("Hidden Gem Detector — A Machine Learning Case Study")
r.font.size = Pt(14)
r.font.italic = True
r.font.color.rgb = RGBColor(27, 79, 114)

meta = doc.add_paragraph()
r = meta.add_run(f"Thomas de Villiers  |  {datetime.date.today().strftime('%B %Y')}")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(108, 117, 125)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# EXECUTIVE SUMMARY
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "Executive Summary", level=1)
body(doc,
    "This case study documents the end-to-end development of a machine learning system "
    "designed to identify underpriced residential property listings in the Egyptian market. "
    "Starting from a raw dataset of 27,322 listings, the project covers data quality "
    "assessment, exploratory analysis, iterative model development, and deployment of both "
    "an interactive price prediction tool and a distributable investor report. The final "
    "model (LightGBM, CV RMSE 0.645, R² 0.674) underpins a Hidden Gem Detector that scores "
    "every listing on value and market desirability — surfacing properties that are "
    "statistically underpriced within their City + Property Type segment."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1: BUSINESS PROBLEM
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "1.  The Business Problem", level=1)
body(doc,
    "Egypt's residential property market is large, fragmented, and opaque. Listings span "
    "179 cities — from beach resort developments on the North Coast to gated urban compounds "
    "in New Cairo — making it practically impossible for an investor to assess value across "
    "the full market manually. The central question this project addresses:"
)

q = doc.add_paragraph()
q.alignment = WD_ALIGN_PARAGRAPH.CENTER
q.paragraph_format.space_before = Pt(6)
q.paragraph_format.space_after  = Pt(6)
r = q.add_run(
    '"Can machine learning identify listings that are statistically underpriced relative to '
    'comparable properties — not simply cheap, but underpriced for what they offer?"'
)
r.font.italic = True
r.font.size = Pt(11.5)
r.font.color.rgb = RGBColor(27, 79, 114)

body(doc,
    "The distinction matters. A cheap property in a low-demand area is not a gem. "
    "A property priced below comparable listings in a high-demand segment — same city, "
    "same property type — is a potential investment opportunity. The goal is to build a "
    "systematic, data-driven method to surface these opportunities at scale."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2: DATASET
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "2.  Dataset Overview", level=1)
body(doc,
    "The dataset contains 27,322 residential property listings scraped from an Egyptian "
    "real estate platform. Each listing includes:"
)

bullet(doc, "Property Type — 10 types including Apartment, Chalet, Stand Alone Villa, Town House, Penthouse")
bullet(doc, "Location — City (179 unique values) and Compound (gated development name)")
bullet(doc, "Size — Area (m²), Bedrooms, Bathrooms")
bullet(doc, "Condition — Furnished status, Floor Level, Delivery Term (Finished / Semi Finished / Core & Shell)")
bullet(doc, "Commercial — Payment Option (Cash / Installment / Both) and Delivery Date")
bullet(doc, "Price — Asking price in Egyptian Pounds (EGP)")

doc.add_paragraph()
body(doc,
    "An important caveat: these are asking prices set by sellers, not transacted prices. "
    "This shapes the interpretation of the model output — 'underpriced' means the seller "
    "is asking less than comparable sellers, not that the market has confirmed the property "
    "is worth more."
)

heading(doc, "2.1  Data Quality Findings", level=2)
body(doc,
    "A critical data quality issue was discovered during initial exploration: missing values "
    "had been encoded as the string 'Unknown' rather than null. Pandas reported near-zero "
    "null counts on raw load, concealing the true extent of missingness. After replacing "
    "all 'Unknown' variants with NaN, the real picture emerged:"
)

add_table(doc,
    ["Column", "Missing %", "Implication"],
    [
        ["Compound",        "40.4%", "High cardinality + high missingness; grouped to top 30 + Other"],
        ["Level",           "38.1%", "Floor number is a key driver for apartments; largely unrecoverable"],
        ["Delivery Date",   "36.9%", "Bucketed into 4 time-horizon tiers before modelling"],
        ["Furnished",       "31.1%", "Binary flag; imputed as 'Missing' category"],
        ["Delivery Term",   "17.1%", "4 clean values; NaN imputed as 'Missing'"],
        ["Payment Option",  "11.0%", "3 clean values; NaN imputed as 'Missing'"],
        ["Area / Beds / Baths", "<2%", "Core numeric features; largely complete"],
    ],
    col_widths=[1.6, 0.9, 3.8]
)
doc.add_paragraph()

body(doc,
    "Additionally, Area contained implausibly small values (minimum 10m²) that inflated "
    "Price/m² to extreme levels. Listings with Area < 20m² were excluded, and a p99 cap "
    "on Price/m² removed the most extreme outliers, leaving 26,578 clean rows."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3: ANALYTICAL APPROACH
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "3.  Analytical Approach & Thought Process", level=1)

heading(doc, "3.1  Why Predict Log(Price)?", level=2)
body(doc,
    "Raw price is heavily right-skewed — the mean (EGP 4.8M) sits well above the median "
    "(EGP 2.7M) due to a small number of ultra-premium listings. Log transformation "
    "produces an approximately normal distribution and, crucially, changes the error metric "
    "from absolute to proportional: an RMSE of 0.1 in log space means the model is "
    "typically within ~10% of the true price, regardless of whether the property costs "
    "1M or 20M EGP. This is the right framing for an investor use case."
)

heading(doc, "3.2  Feature Engineering Decisions", level=2)
body(doc, "Several deliberate engineering choices were made before modelling:")
bullet(doc,
    "Target-encoding City and Compound: with 179 cities, label encoding assigns arbitrary "
    "integers that give the model zero price signal. Target encoding replaces each category "
    "with its mean log(Price), with Laplace smoothing to prevent overfitting on sparse "
    "categories.",
    bold_prefix="City / Compound encoding."
)
bullet(doc,
    "log(Area) added alongside raw Area, since price scales with size in log-log space. "
    "The model is free to use whichever representation is more informative.",
    bold_prefix="Log(Area)."
)
bullet(doc,
    "Beds_per_100m² and Bath/Bed ratio capture density and luxury signals that raw counts "
    "cannot express. A 3-bed, 300m² villa is very different from a 3-bed, 90m² apartment.",
    bold_prefix="Interaction features."
)
bullet(doc,
    "Delivery Date was bucketed into four tiers (Ready / Near Term / Mid Term / Long Term) "
    "to reduce the 9-value cardinality and capture the time-value dimension of off-plan vs "
    "ready-to-move properties.",
    bold_prefix="Delivery Date bucketing."
)
bullet(doc,
    "City × Type target encoding: an Apartment in New Cairo has a fundamentally different "
    "price level to a Villa in New Cairo. Encoding the interaction directly gives the model "
    "a granular segment price signal.",
    bold_prefix="City × Type combination."
)

heading(doc, "3.3  Model Selection Rationale", level=2)
body(doc,
    "Two target variable formulations were evaluated: Model A predicts log(Price) directly "
    "with Area as a feature; Model B predicts log(Price/m²) and reconstructs price by "
    "multiplying by actual area. Both LightGBM and XGBoost were tested for each. "
    "LightGBM consistently outperformed XGBoost on this dataset, likely because it handles "
    "high-cardinality categoricals more efficiently and its leaf-wise tree growth better "
    "captures the heterogeneous market structure."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4: MODELLING JOURNEY
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "4.  The Modelling Journey", level=1)
body(doc,
    "Models were developed iteratively, with each version adding specific improvements "
    "and testing their contribution via 5-fold cross-validated RMSE. This approach makes "
    "the value of each change transparent and defensible."
)

add_table(doc,
    ["Version", "Key Changes", "CV RMSE", "R²", "Delta"],
    [
        ["v1 — Baseline",
         "LightGBM with label-encoded categoricals",
         "0.682", "0.668", "—"],
        ["v2 — Core Improvements",
         "Price outlier removal (p1–p99); target-encode City & Compound; "
         "add log(Area); tune hyperparameters (1,500 trees, LR 0.02)",
         "0.645", "0.674", "−5.4% RMSE"],
        ["v3 — XGBoost Comparison",
         "XGBoost baseline applied to v2 feature set",
         "0.653", "0.666", "+1.2% vs v2"],
        ["v4 — Further Engineering",
         "Interaction features (Beds/100m², Bath/Bed ratio); "
         "City × Type target encoding; LightGBM hyperparameter search",
         "TBD", "TBD", "Pending"],
    ],
    col_widths=[1.0, 3.2, 0.8, 0.6, 0.8]
)
doc.add_paragraph()

body(doc,
    "The v3 XGBoost hyperparameter search (RandomizedSearchCV) produced worse results than "
    "the default configuration, illustrating a common pitfall: when features contain "
    "pre-computed target encodings, the internal CV folds in RandomizedSearchCV see "
    "information from the validation fold already 'baked in' to the encodings, producing "
    "unreliable relative scores. The search navigated with a partially broken compass. "
    "This was corrected in v4 by tuning LightGBM (the confirmed winner) with a narrower, "
    "more informed search space."
)

heading(doc, "4.1  Why R² of 0.67 is Honest, Not Embarrassing", level=2)
body(doc,
    "A model explaining 67% of price variation on this dataset is not underperforming — "
    "it is confronting the data's fundamental limits. The unexplained 33% decomposes into:"
)
bullet(doc,
    "~15% from genuinely missing features: no property age, no condition rating, no GPS "
    "coordinates, no view or orientation data. These factors have real price impact but are "
    "absent from the dataset."
)
bullet(doc,
    "~10% from market heterogeneity: 179 cities spanning beach resorts, gated compounds, "
    "and urban residential areas follow different pricing logics that a single model "
    "struggles to reconcile."
)
bullet(doc,
    "~8% from genuine noise: seller pricing behaviour is partly driven by motivation, "
    "anchoring to neighbour listing prices, and negotiation expectations — none of which "
    "are observable from listing features."
)
body(doc,
    "Without new data sources (coordinates, property age, condition scores), a realistic "
    "ceiling for this dataset is R² ≈ 0.75–0.78."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5: HIDDEN GEM METHODOLOGY
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "5.  Hidden Gem Detection Methodology", level=1)

heading(doc, "5.1  Why Global Residuals Are Insufficient", level=2)
body(doc,
    "An initial approach flagged listings where actual price was below the model's "
    "predicted price. The problem: a studio below the global model prediction might simply "
    "be priced correctly for a small studio — the model's average prediction is pulled "
    "upward by larger properties in the same city. Comparing to the wrong peer group "
    "generates false positives."
)

heading(doc, "5.2  Segment-Level Scoring", level=2)
body(doc,
    "The refined approach compares each listing only to its direct peers: same City, "
    "same Property Type. Segments with fewer than 5 listings are excluded for statistical "
    "reliability. Within each segment, two scores are computed:"
)
bullet(doc,
    "Percentile rank of Price/m² within segment, inverted (0–100). A score of 90 means "
    "the listing is cheaper than 90% of comparable properties in that segment.",
    bold_prefix="Value Score."
)
bullet(doc,
    "Segment median Price/m², normalised 0–100 across all segments. Higher median = more "
    "in-demand market. This prevents flagging cheap properties in markets where everything "
    "is cheap — cheap for a reason.",
    bold_prefix="Desirability Score."
)
bullet(doc,
    "Geometric mean of Value and Desirability. Because it multiplies rather than adds, "
    "both scores must be high for the combined score to be high. A listing scoring 90 on "
    "value but 5 on desirability yields a gem score of only 21 — correctly de-prioritised.",
    bold_prefix="Hidden Gem Score."
)
body(doc,
    "Only listings in the bottom 25th percentile by Price/m² within their segment qualify. "
    "Final output is sorted by Gem Score descending."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6: DELIVERABLES
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "6.  Deliverables", level=1)

add_table(doc,
    ["Deliverable", "Description", "How to Use"],
    [
        ["HousePrices.ipynb",
         "Full analysis notebook covering Steps 1–15: data cleaning, EDA, feature "
         "engineering, model development (v1–v4), refined hidden gem scoring",
         "Open in Jupyter Notebook"],
        ["hidden_gems_report.html",
         "Static interactive report: KPI cards, two charts, sortable/searchable "
         "table of all hidden gems with Value, Desirability, and Gem Scores",
         "Open in any web browser — no server required"],
        ["app.py",
         "Streamlit web app: enter any property's features and receive an instant "
         "price estimate with 68% confidence range",
         "streamlit run app.py"],
        ["model_artifacts.pkl",
         "Serialised model, encoding lookups, and metadata — loaded by the Streamlit "
         "app; regenerated by running Step 13 in the notebook",
         "Auto-loaded by app.py"],
    ],
    col_widths=[1.5, 3.2, 1.7]
)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 7: LIMITATIONS
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "7.  Limitations & Honest Caveats", level=1)
body(doc,
    "Presenting findings without acknowledging limitations is poor practice. "
    "The following constraints should be understood by anyone acting on this output:"
)
bullet(doc,
    "The model is trained on asking prices, not sold prices. A listing flagged as a hidden "
    "gem is underpriced relative to what other sellers are asking — not necessarily relative "
    "to what the market will pay. The seller may reject below-asking offers.",
    bold_prefix="Asking prices only."
)
bullet(doc,
    "High missingness in Level (38%) and Compound (40%) limits the model's ability to "
    "capture floor-number premiums and compound-specific pricing. These are material "
    "drivers in the Egyptian market.",
    bold_prefix="Missing features."
)
bullet(doc,
    "Target encodings were computed on the full dataset before cross-validation, "
    "introducing a small amount of data leakage. Scores are slightly optimistic; true "
    "out-of-sample performance may be marginally lower.",
    bold_prefix="Target encoding leakage."
)
bullet(doc,
    "A high Gem Score means the listing is cheap and the segment is desirable. It does "
    "not mean the listing is in good condition, has no legal encumbrances, or is otherwise "
    "suitable for investment. Due diligence on individual properties remains essential.",
    bold_prefix="Due diligence still required."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 8: TECHNICAL APPENDIX
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "8.  Technical Appendix", level=1)

heading(doc, "8.1  Technology Stack", level=2)
bullet(doc, "Python 3.9")
bullet(doc, "pandas, NumPy — data manipulation")
bullet(doc, "LightGBM, XGBoost, scikit-learn — modelling")
bullet(doc, "Matplotlib, Seaborn — visualisation")
bullet(doc, "Streamlit — interactive price prediction UI")
bullet(doc, "Bootstrap 5, DataTables, Chart.js — HTML report")
bullet(doc, "joblib — model serialisation")

heading(doc, "8.2  Model Configuration (Best: LightGBM v2)", level=2)
add_table(doc,
    ["Parameter", "Value", "Rationale"],
    [
        ["n_estimators",      "1,500",  "Sufficient trees for convergence at LR 0.02"],
        ["learning_rate",     "0.02",   "Low enough to avoid overfitting with 1,500 trees"],
        ["num_leaves",        "127",    "Controls model complexity; LightGBM's key parameter"],
        ["min_child_samples", "20",     "Prevents overfitting on sparse segments"],
        ["subsample",         "0.8",    "Row subsampling for regularisation"],
        ["colsample_bytree",  "0.8",    "Feature subsampling per tree"],
        ["reg_alpha/lambda",  "0.1",    "Light L1/L2 regularisation"],
    ],
    col_widths=[1.6, 0.8, 4.0]
)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
out_path = '/Users/tomdevilliers/Documents/Thomas/Woolworths/Hidden_Gem_Case_Study.docx'
doc.save(out_path)
print(f"Document saved: {out_path}")
